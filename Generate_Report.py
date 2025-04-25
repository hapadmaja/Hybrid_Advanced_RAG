import time
start = time.time()
from llama_index.llms.openai import OpenAI
import streamlit as st
from llama_index.embeddings.openai import OpenAIEmbedding
from llama_index.core import VectorStoreIndex,Settings
import asyncio

import pickle

from llama_index.core.tools import QueryEngineTool
import os
from docx import Document
from llama_index.core.workflow import (
    step,
    Event,
    Context,
    StartEvent,
    StopEvent,
    Workflow,
)
from llama_index.core.agent import FunctionCallingAgent

os.environ["LLAMA_CLOUD_API_KEY"] = "YOUR_LLAMA_KEY"
# Using OpenAI API for embeddings/llms
os.environ["OPENAI_API_KEY"] =  "YOUR_OPRN_AI_KEY"
st.title("📄  Report Generation Agent")

end = time.time()

print(f"The time required for loading libraries :{end - start}")

if "initialized" not in st.session_state:
    st.session_state.initialized = False

if not st.session_state.initialized:
    start = time.time()
    embed_model = OpenAIEmbedding(model="text-embedding-3-small")
    llm = OpenAI(model="gpt-4o-mini")

    Settings.llm = llm
    Settings.embed_model = embed_model


    st.write("Loading the files")
    print("Loading the files")
    # to convert pickle to documents form   
    pickle_file_path = "financial_documents.pkl"
    with open(pickle_file_path, "rb") as f:
        p_documents = pickle.load(f)
    print(f"Loaded {len(p_documents)} documents from pickle!")

    st.write("Creating Index")
    print("Creating Index")
    index = VectorStoreIndex.from_documents(p_documents)
    st.write("Initialising Query Engine")
    print("Initialising Query Engine")
    query_engine = index.as_query_engine(similarity_top_k=10)
    budget_tool = QueryEngineTool.from_defaults(
    query_engine,
    name="Indian_economic_survey_2425",
    description="A RAG engine with extremely detailed information about the 2024-25 Indian economic survey.",
    )
    
    st.session_state.embed_model=embed_model
    st.session_state.llm=llm
    st.session_state.p_documents=p_documents
    st.session_state. index=index
    st.session_state.query_engine = query_engine
    st.session_state.budget_tool=budget_tool
    st.session_state.initialized = True  # Mark as initialized
    st.write("Initialization complete! ✅")

    end = time.time()
    print(f"The time required for initialisation :{end - start}")

embed_model=st.session_state.embed_model
llm=st.session_state.llm
p_documents=st.session_state.p_documents
index=st.session_state. index
query_engine=st.session_state.query_engine
budget_tool=st.session_state.budget_tool

class OutlineEvent(Event):
        outline: str
    
class QuestionEvent(Event):
    question: str


class AnswerEvent(Event):
    question: str
    answer: str


class ReviewEvent(Event):
    report: str


class ProgressEvent(Event):
    progress: str

    
class DocumentResearchAgent(Workflow):
    # get the initial request and create an outline of the report post knowing nothing about the topic
    @step()
    async def formulate_plan(
        self, ctx: Context, ev: StartEvent
    ) -> OutlineEvent:
        query = ev.query
        await ctx.set("original_query", query)
        await ctx.set("tools", ev.tools)
        st.write("Formulating Outline")
        prompt = f"""You are an expert at writing blog post. You have been given a topic to write
        a blog post about. Plan an outline for the report ; it should be detailed and specific.
        Another agent will formulate questions to find the facts necessary to fulfill the outline.
        The topic is: {query}"""

        response = await Settings.llm.acomplete(prompt)
        text=str(response)
        modified_text = text.replace("Blog Post Outline: ", "Report on ")
        ctx.write_event_to_stream(
            ProgressEvent(progress="Outline:\n" + modified_text)
        )
        st.write("Outline", modified_text)
        return OutlineEvent(outline=modified_text)

    # formulate some questions based on the outline
    @step()
    async def formulate_questions(
        self, ctx: Context, ev: OutlineEvent
    ) -> QuestionEvent:
        outline = ev.outline
        await ctx.set("outline", outline)
        st.write("Formulating Questions")
        prompt = f"""You are an expert at formulating research questions. You have been given an outline
        for a blog post. Formulate a series of simple questions that will get you the facts necessary
        to fulfill the outline. You cannot assume any existing knowledge; you must ask at least one
        question for every bullet point in the outline. Avoid complex or multi-part questions; break
        them down into a series of simple questions. Your output should be a list of questions, each
        on a new line. Do not include headers or categories or any preamble or explanation; just a
        list of questions. For speed of response, limit yourself to 8 questions. The outline is: {outline}"""

        response = await Settings.llm.acomplete(prompt)

        questions = str(response).split("\n")
        questions = [x for x in questions if x]

        ctx.write_event_to_stream(
            ProgressEvent(
                progress="Formulated questions:\n" + "\n".join(questions)
            )
        )

        await ctx.set("num_questions", len(questions))

        ctx.write_event_to_stream(
            ProgressEvent(progress="Questions:\n" + "\n".join(questions))
        )
        st.write("Questions : ")
        for question in questions:
            ctx.send_event(QuestionEvent(question=question))
            st.write(question)

    # answer each question in turn
    @step()
    async def answer_question(
        self, ctx: Context, ev: QuestionEvent
    ) -> AnswerEvent:
        question = ev.question
        if (
            not question
            or question.isspace()
            or question == ""
            or question is None
        ):
            ctx.write_event_to_stream(
                ProgressEvent(progress=f"Skipping empty question.")
            )  # Log skipping empty question
            return None
        agent = FunctionCallingAgent.from_tools(
            await ctx.get("tools"),
            verbose=True,
        )
        response = await agent.aquery(question)

        ctx.write_event_to_stream(
            ProgressEvent(
                progress=f"To question '{question}' the agent answered: {response}"
            )
        )

        return AnswerEvent(question=question, answer=str(response))

    # given all the answers to all the questions and the outline, write the blog poost
    @step()
    async def write_report(self, ctx: Context, ev: AnswerEvent) -> ReviewEvent:
        # wait until we receive as many answers as there are questions
        num_questions = await ctx.get("num_questions")
        results = ctx.collect_events(ev, [AnswerEvent] * num_questions)
        if results is None:
            return None

        # maintain a list of all questions and answers no matter how many times this step is called
        try:
            previous_questions = await ctx.get("previous_questions")
        except:
            previous_questions = []
        previous_questions.extend(results)
        await ctx.set("previous_questions", previous_questions)

        prompt = f"""You are an expert at writing blog posts. You are given an outline of a blog post
        and a series of questions and answers that should provide all the data you need to write the
        blog post. Compose the blog post according to the outline, using only the data given in the
        answers. The outline is in  and the questions and answers are in  and
        .
        {await ctx.get('outline')}"""

        for result in previous_questions:
            prompt += f"{result.question}\n{result.answer}\n"

        ctx.write_event_to_stream(
            ProgressEvent(progress="Writing report with prompt:\n" + prompt)
        )
        st.write("Writing Report")
        report = await Settings.llm.acomplete(prompt)

        return ReviewEvent(report=str(report))

    # review the report. If it still needs work, formulate some more questions.
    @step
    async def review_report(
        self, ctx: Context, ev: ReviewEvent
    ) -> StopEvent | QuestionEvent:
        # we re-review a maximum of 3 times
        try:
            num_reviews = await ctx.get("num_reviews")
        except:
            num_reviews = 1
        num_reviews += 1
        await ctx.set("num_reviews", num_reviews)

        report = ev.report
        st.write("Reviewing the Report")
        prompt = f"""You are an expert reviewer of blog post. You are given an original query,
        and a blog post that was written to satisfy that query. Review the blog post and determine
        if it adequately answers the query and contains enough detail. If it doesn't, come up with
        a set of questions that will get you the facts necessary to expand the blog. Another
        agent will answer those questions. Your response should just be a list of questions, one
        per line, without any preamble or explanation. For speed, generate a maximum of 4 questions.
        The original query is: '{await ctx.get('original_query')}'.
        The blog is: {report}.
        If the blog is fine, return just the string 'OKAY'."""

        response = await Settings.llm.acomplete(prompt)

        if response == "OKAY" or await ctx.get("num_reviews") >= 3:
            ctx.write_event_to_stream(
                ProgressEvent(progress="Report content is fine")
            )
            return StopEvent(result=report)
        else:
            questions = str(response).split("\n")
            await ctx.set("num_questions", len(questions))
            ctx.write_event_to_stream(
                ProgressEvent(progress="Formulated some more questions")
            )
            for question in questions:
                ctx.send_event(QuestionEvent(question=question))

async def run_agent(query):
    
    agent = DocumentResearchAgent(timeout=600, verbose=True)
    handler = agent.run(query=query, tools=[budget_tool])

    progress_messages = []

    async for ev in handler.stream_events():
        if isinstance(ev, ProgressEvent):
            progress_messages.append(ev.progress)
            #st.session_state["progress"] = "\n".join(progress_messages)
            st.session_state["status"] = "running"
            print(ev.progress)
            #st.write(ev.progress)
            #st.rerun()

    final_result = await handler
    st.session_state["final_result"] = final_result
    st.session_state["status"] = "done"

    # Save the result to a Word document
    doc = Document()
    doc.add_heading(f"Report for {query}", level=1)
    doc.add_paragraph(final_result)
    doc_path = "Report.docx"
    doc.save(doc_path)

    st.session_state["doc_path"] = doc_path


# Streamlit UI
FOLDER_PATH = "./economic_survey"  # Change this to your actual folder path

# Function to fetch PDF document names from the folder
def get_loaded_documents(folder_path):
    try:
        files = [f for f in os.listdir(folder_path) if f.endswith(".pdf")]
        return files
    except Exception as e:
        return [f"Error accessing folder: {e}"]
st.markdown(
    """
    <style>
        .title {
            color: #004aad;
            font-size: 28px;
            font-weight: bold;
            text-align: center;
        }
        .sidebar-text {
            font-size: 16px;
            font-weight: bold;
            color: #333;
        }
    </style>
    """,
    unsafe_allow_html=True,
)
st.sidebar.title("📂 Documents Loaded")

# Fetch PDF document names
documents = get_loaded_documents(FOLDER_PATH)

if documents:
    for doc in documents:
        st.sidebar.markdown(f'<p class="sidebar-text">📄 {doc}</p>', unsafe_allow_html=True)
else:
    st.sidebar.warning("No PDFs found in the folder.")

query = st.text_area("Enter your research query", "Give me a comprehensive analysis on Indian Economic Survey 2024-25")

if st.button("Generate Report"):
    if query.strip():
        st.session_state["progress"] = ""
        st.session_state["final_result"] = None
        st.session_state["status"] = "running"
        #asyncio.create_task(run_agent(query))  
        asyncio.run(run_agent(query))# Run asynchronously in the background
    else:
        st.warning("Please enter a query.")

# Display progress
if st.session_state.get("status") == "running":
    st.text_area("Processing Updates", st.session_state["progress"], height=200)

# Display final result and download option
if st.session_state.get("status") == "done":
    st.success("Report generation complete!")
    st.text_area("Final Report", st.session_state["final_result"], height=200)

    with open(st.session_state["doc_path"], "rb") as file:
        st.download_button("Download Report", file, "Indian_Economic_Survey_Analysis.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

st.sidebar.title("About")
st.sidebar.info("This is an report generation system designed for financial document analysis.")
